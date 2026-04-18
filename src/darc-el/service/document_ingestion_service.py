from __future__ import annotations

import io
import mimetypes
from pathlib import Path
from typing import Any, ClassVar

from core.document_ingestion import DocumentChunk, DocumentIngestionRecord
from pydantic import BaseModel


class UnsupportedDocumentTypeError(ValueError):
    pass


class DocumentIngestionService(BaseModel):
    supported_types: ClassVar[dict[str, str]] = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "text",
    }

    def ingest_upload(self, file_name: str, content_type: str | None, data: bytes) -> DocumentIngestionRecord:
        detected_type = self._detect_type(file_name, content_type)
        metadata: dict[str, Any]
        text: str
        parser_name: str

        if detected_type == "pdf":
            metadata, text, parser_name = self._parse_pdf(data)
        elif detected_type == "docx":
            metadata, text, parser_name = self._parse_docx(data)
        elif detected_type == "text":
            metadata, text, parser_name = self._parse_text(file_name, data)
        else:  # pragma: no cover - guarded by supported_types
            raise UnsupportedDocumentTypeError(f"Unsupported document type: {content_type or file_name}")

        chunks = self._chunk_text(text)
        return DocumentIngestionRecord(
            file_name=file_name,
            content_type=content_type or mimetypes.guess_type(file_name)[0] or "application/octet-stream",
            source_type=detected_type,
            metadata=metadata,
            text=text,
            chunks=chunks,
            parser_name=parser_name,
        )

    def _detect_type(self, file_name: str, content_type: str | None) -> str:
        if content_type in self.supported_types:
            return self.supported_types[content_type]

        suffix = Path(file_name).suffix.lower()
        if suffix == ".pdf":
            return "pdf"
        if suffix == ".docx":
            return "docx"
        if suffix in {".txt", ".text"}:
            return "text"

        raise UnsupportedDocumentTypeError(f"Unsupported document type: {content_type or file_name}")

    def _parse_pdf(self, data: bytes) -> tuple[dict[str, Any], str, str]:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
        text_parts: list[str] = []
        for page in reader.pages:
            page_text = page.extract_text() or ""
            if page_text:
                text_parts.append(page_text)

        metadata = {
            "page_count": len(reader.pages),
            "pdf_metadata": self._stringify_metadata(reader.metadata or {}),
        }
        return metadata, "\n".join(text_parts).strip(), "pypdf"

    def _parse_docx(self, data: bytes) -> tuple[dict[str, Any], str, str]:
        from docx import Document

        document = Document(io.BytesIO(data))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        metadata = {
            "paragraph_count": len(document.paragraphs),
            "core_properties": self._docx_properties(document),
        }
        return metadata, "\n".join(paragraphs).strip(), "python-docx"

    def _parse_text(self, file_name: str, data: bytes) -> tuple[dict[str, Any], str, str]:
        text = data.decode("utf-8", errors="replace")
        return {"encoding": "utf-8", "file_extension": Path(file_name).suffix.lower()}, text, "utf-8-text"

    def _chunk_text(self, text: str, chunk_size: int = 1200, overlap: int = 150) -> list[DocumentChunk]:
        normalized = text.strip()
        if not normalized:
            return []

        chunks: list[DocumentChunk] = []
        start = 0
        index = 0
        length = len(normalized)
        while start < length:
            end = min(length, start + chunk_size)
            chunk_text = normalized[start:end].strip()
            if chunk_text:
                chunks.append(DocumentChunk(index=index, text=chunk_text))
                index += 1
            if end >= length:
                break
            start = max(end - overlap, start + 1)
        return chunks

    def _stringify_metadata(self, metadata: dict[str, Any]) -> dict[str, Any]:
        return {key: (value if isinstance(value, (str, int, float, bool)) or value is None else str(value)) for key, value in metadata.items()}

    def _docx_properties(self, document: Any) -> dict[str, Any]:
        properties = getattr(document, "core_properties", None)
        if properties is None:
            return {}
        return {
            "author": getattr(properties, "author", None),
            "title": getattr(properties, "title", None),
            "subject": getattr(properties, "subject", None),
            "keywords": getattr(properties, "keywords", None),
            "created": getattr(properties, "created", None).isoformat() if getattr(properties, "created", None) else None,
            "modified": getattr(properties, "modified", None).isoformat() if getattr(properties, "modified", None) else None,
        }

    def to_payload(self, record: DocumentIngestionRecord) -> dict[str, Any]:
        return {
            "file_name": record.file_name,
            "content_type": record.content_type,
            "source_type": record.source_type,
            "metadata": record.metadata,
            "text": record.text,
            "chunks": [chunk.model_dump() for chunk in record.chunks],
            "parser_name": record.parser_name,
        }
